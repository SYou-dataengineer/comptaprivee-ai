"""Conversions locales de documents pour ComptaPrivée AI."""

from __future__ import annotations

import csv
import sys
from dataclasses import dataclass
from pathlib import Path

import fitz

from .pdf_table_extractor import extraire_tableaux_pdf


EXTENSIONS_WORD = {".doc", ".docx"}
EXTENSIONS_EXCEL = {".xls", ".xlsx", ".xlsm"}
EXTENSIONS_IMAGES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff"}

CONVERSIONS_SUPPORTEES = (
    "Word → PDF",
    "Excel → PDF",
    "Excel → CSV",
    "CSV → Excel",
    "Images → PDF",
    "PDFs → PDF",
    "PDF → Images (PNG)",
    "PDF → Images (JPG)",
    "PDF → CSV",
    "PDF → Excel",
    "PDF → Word",
)


class ErreurConversion(RuntimeError):
    """Erreur contrôlée lors d'une conversion locale."""


@dataclass(frozen=True)
class ResultatConversion:
    """Résultat d'une conversion locale."""

    source: Path
    destination: Path
    type_conversion: str


def _verifier_source(
    source: str | Path,
    extensions: set[str],
) -> Path:
    chemin = Path(source)

    if not chemin.exists():
        raise FileNotFoundError(
            f"Fichier source introuvable : {chemin}"
        )

    if chemin.suffix.lower() not in extensions:
        attendues = ", ".join(sorted(extensions))
        raise ValueError(
            f"Format source non pris en charge. Formats attendus : {attendues}"
        )

    return chemin


def _preparer_destination(
    source: Path,
    destination: str | Path | None,
    extension: str,
) -> Path:
    if destination is None:
        chemin = source.with_suffix(extension)
    else:
        chemin = Path(destination)

    if chemin.suffix.lower() != extension:
        raise ValueError(
            f"Le fichier de destination doit avoir l'extension {extension}."
        )

    chemin.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    return chemin


def word_vers_pdf(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Convertit Word en PDF via Microsoft Word installé sur Windows."""
    source_path = _verifier_source(
        source,
        EXTENSIONS_WORD,
    )
    destination_path = _preparer_destination(
        source_path,
        destination,
        ".pdf",
    )

    if sys.platform != "win32":
        raise ErreurConversion(
            "Word → PDF nécessite Windows et Microsoft Word installé."
        )

    try:
        import win32com.client  # type: ignore
    except ImportError as erreur:
        raise ErreurConversion(
            "Word → PDF nécessite le module pywin32. "
            "Installez les dépendances du projet."
        ) from erreur

    word = None

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        document = word.Documents.Open(
            str(source_path.resolve()),
            ReadOnly=True,
        )
        try:
            document.ExportAsFixedFormat(
                str(destination_path.resolve()),
                17,
            )
        finally:
            document.Close(False)

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de convertir le document Word : {erreur}"
        ) from erreur

    finally:
        if word is not None:
            word.Quit()

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="Word → PDF",
    )


def excel_vers_pdf(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Convertit Excel en PDF, une feuille visible = une page PDF."""
    source_path = _verifier_source(
        source,
        EXTENSIONS_EXCEL,
    )
    destination_path = _preparer_destination(
        source_path,
        destination,
        ".pdf",
    )

    if sys.platform != "win32":
        raise ErreurConversion(
            "Excel → PDF nécessite Windows et Microsoft Excel installé."
        )

    try:
        import win32com.client  # type: ignore
    except ImportError as erreur:
        raise ErreurConversion(
            "Excel → PDF nécessite le module pywin32. "
            "Installez les dépendances du projet."
        ) from erreur

    excel = None

    try:
        excel = win32com.client.DispatchEx("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        excel.ScreenUpdating = False
        excel.EnableEvents = False

        classeur = excel.Workbooks.Open(
            str(source_path.resolve()),
            ReadOnly=True,
            UpdateLinks=0,
        )

        try:
            feuilles_exportees = 0

            for feuille in classeur.Worksheets:
                # Ignorer les feuilles masquées.
                if feuille.Visible != -1:
                    continue

                # Trouver la vraie zone de contenu.
                # UsedRange peut être artificiellement énorme à cause
                # d'anciens formats, ce qui produit un PDF minuscule.
                premiere_cellule = feuille.Cells.Find(
                    What="*",
                    After=feuille.Cells(1, 1),
                    LookIn=-4163,      # xlFormulas
                    LookAt=2,         # xlPart
                    SearchOrder=1,    # xlByRows
                    SearchDirection=1 # xlNext
                )

                derniere_ligne = feuille.Cells.Find(
                    What="*",
                    After=feuille.Cells(1, 1),
                    LookIn=-4163,
                    LookAt=2,
                    SearchOrder=1,     # xlByRows
                    SearchDirection=2, # xlPrevious
                )

                derniere_colonne = feuille.Cells.Find(
                    What="*",
                    After=feuille.Cells(1, 1),
                    LookIn=-4163,
                    LookAt=2,
                    SearchOrder=2,     # xlByColumns
                    SearchDirection=2, # xlPrevious
                )

                if (
                    premiere_cellule is None
                    or derniere_ligne is None
                    or derniere_colonne is None
                ):
                    continue

                premiere_ligne = premiere_cellule.Row

                premiere_colonne_cellule = feuille.Cells.Find(
                    What="*",
                    After=feuille.Cells(1, 1),
                    LookIn=-4163,
                    LookAt=2,
                    SearchOrder=2,    # xlByColumns
                    SearchDirection=1 # xlNext
                )

                if premiere_colonne_cellule is None:
                    continue

                premiere_colonne = premiere_colonne_cellule.Column
                derniere_ligne_no = derniere_ligne.Row
                derniere_colonne_no = derniere_colonne.Column

                zone = feuille.Range(
                    feuille.Cells(
                        premiere_ligne,
                        premiere_colonne,
                    ),
                    feuille.Cells(
                        derniere_ligne_no,
                        derniere_colonne_no,
                    ),
                )

                try:
                    excel.PrintCommunication = False
                except Exception:
                    pass

                try:
                    feuille.PageSetup.PrintArea = zone.Address

                    # Une feuille Excel visible = exactement une page PDF.
                    feuille.PageSetup.Zoom = False
                    feuille.PageSetup.FitToPagesWide = 1
                    feuille.PageSetup.FitToPagesTall = 1

                    largeur_points = float(zone.Width)
                    hauteur_points = float(zone.Height)

                    # Orientation réellement basée sur les dimensions
                    # visuelles du contenu, pas seulement le nombre de cellules.
                    feuille.PageSetup.Orientation = (
                        2 if largeur_points > hauteur_points else 1
                    )

                    # Papier Letter : format courant au Canada.
                    feuille.PageSetup.PaperSize = 1

                    # Marges compactes mais propres.
                    feuille.PageSetup.LeftMargin = excel.InchesToPoints(0.20)
                    feuille.PageSetup.RightMargin = excel.InchesToPoints(0.20)
                    feuille.PageSetup.TopMargin = excel.InchesToPoints(0.25)
                    feuille.PageSetup.BottomMargin = excel.InchesToPoints(0.25)
                    feuille.PageSetup.HeaderMargin = 0
                    feuille.PageSetup.FooterMargin = 0

                    feuille.PageSetup.CenterHorizontally = True
                    feuille.PageSetup.CenterVertically = True

                    # Évite des pages supplémentaires dues aux titres répétés.
                    feuille.PageSetup.PrintTitleRows = ""
                    feuille.PageSetup.PrintTitleColumns = ""

                    feuilles_exportees += 1

                finally:
                    try:
                        excel.PrintCommunication = True
                    except Exception:
                        pass

            if feuilles_exportees == 0:
                raise ErreurConversion(
                    "Aucune feuille Excel visible contenant des données "
                    "n'a été trouvée."
                )

            # Toutes les feuilles visibles configurées sont exportées :
            # chacune tient sur une page PDF.
            classeur.ExportAsFixedFormat(
                0,
                str(destination_path.resolve()),
                0,
                True,
                False,
            )

        finally:
            classeur.Close(False)

    except ErreurConversion:
        raise

    except Exception as erreur:
        message = str(erreur)

        if (
            "Document not saved" in message
            or "Document non enregistré" in message
        ):
            raise ErreurConversion(
                "Impossible d'enregistrer le PDF. "
                "Fermez le PDF de destination s'il est déjà ouvert, "
                "ou choisissez un autre nom de fichier."
            ) from erreur

        raise ErreurConversion(
            f"Impossible de convertir le classeur Excel : {erreur}"
        ) from erreur

    finally:
        if excel is not None:
            try:
                excel.ScreenUpdating = True
                excel.EnableEvents = True
            except Exception:
                pass
            excel.Quit()

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="Excel → PDF",
    )


def excel_vers_csv(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Convertit la feuille active d'un fichier Excel en CSV UTF-8."""
    source_path = _verifier_source(
        source,
        EXTENSIONS_EXCEL,
    )
    destination_path = _preparer_destination(
        source_path,
        destination,
        ".csv",
    )

    try:
        from openpyxl import load_workbook
    except ImportError as erreur:
        raise ErreurConversion(
            "Excel → CSV nécessite openpyxl. "
            "Installez les dépendances du projet."
        ) from erreur

    try:
        classeur = load_workbook(
            source_path,
            data_only=True,
            read_only=False,
        )
        feuille = classeur.active

        # Si la feuille contient un vrai tableau Excel, on exporte
        # uniquement ce tableau. Cela évite les notes ou contenus annexes
        # placés plus bas dans la feuille.
        plage = None

        if feuille.tables:
            premier_tableau = next(iter(feuille.tables.values()))
            plage = premier_tableau.ref

        if plage:
            cellules = feuille[plage]
        else:
            # Sinon, on détecte la zone réellement utilisée en ignorant
            # les lignes et colonnes complètement vides autour des données.
            lignes_non_vides = []
            for ligne in feuille.iter_rows():
                if any(cellule.value is not None for cellule in ligne):
                    lignes_non_vides.append(ligne)

            if not lignes_non_vides:
                raise ErreurConversion(
                    "La feuille Excel ne contient aucune donnée à exporter."
                )

            min_row = min(cellule.row for ligne in lignes_non_vides for cellule in ligne if cellule.value is not None)
            max_row = max(cellule.row for ligne in lignes_non_vides for cellule in ligne if cellule.value is not None)
            min_col = min(cellule.column for ligne in lignes_non_vides for cellule in ligne if cellule.value is not None)
            max_col = max(cellule.column for ligne in lignes_non_vides for cellule in ligne if cellule.value is not None)

            cellules = feuille.iter_rows(
                min_row=min_row,
                max_row=max_row,
                min_col=min_col,
                max_col=max_col,
            )

        with destination_path.open(
            "w",
            encoding="utf-8-sig",
            newline="",
        ) as fichier_csv:
            writer = csv.writer(fichier_csv)

            for ligne in cellules:
                valeurs = []

                for cellule in ligne:
                    valeur = cellule.value

                    if valeur is None:
                        valeurs.append("")
                    elif hasattr(valeur, "strftime"):
                        valeurs.append(
                            valeur.strftime("%Y-%m-%d")
                        )
                    else:
                        valeurs.append(valeur)

                writer.writerow(valeurs)

        classeur.close()

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de convertir Excel en CSV : {erreur}"
        ) from erreur

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="Excel → CSV",
    )


def image_vers_pdf(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Convertit une image prise en charge en PDF local."""
    source_path = _verifier_source(
        source,
        EXTENSIONS_IMAGES,
    )
    destination_path = _preparer_destination(
        source_path,
        destination,
        ".pdf",
    )

    document = fitz.open()

    try:
        image = fitz.open(source_path)
        try:
            pdf_bytes = image.convert_to_pdf()
        finally:
            image.close()

        pdf_image = fitz.open(
            "pdf",
            pdf_bytes,
        )
        try:
            document.insert_pdf(pdf_image)
        finally:
            pdf_image.close()

        document.save(destination_path)

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de convertir l'image en PDF : {erreur}"
        ) from erreur

    finally:
        document.close()

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="Image → PDF",
    )


def images_vers_pdf(
    sources: list[str | Path] | tuple[str | Path, ...],
    destination: str | Path | None = None,
) -> ResultatConversion:
    # Regroupe plusieurs images dans un seul PDF, dans l'ordre fourni.
    if not sources:
        raise ValueError(
            "Sélectionnez au moins une image à convertir."
        )

    chemins = [
        _verifier_source(source, EXTENSIONS_IMAGES)
        for source in sources
    ]

    premiere_source = chemins[0]

    if destination is None:
        destination_path = premiere_source.with_name(
            f"{premiere_source.stem}_images.pdf"
        )
    else:
        destination_path = Path(destination)

    if destination_path.suffix.lower() != ".pdf":
        raise ValueError(
            "Le fichier de destination doit avoir l'extension .pdf."
        )

    destination_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = fitz.open()

    try:
        for chemin_image in chemins:
            image = fitz.open(chemin_image)

            try:
                pdf_bytes = image.convert_to_pdf()
            finally:
                image.close()

            pdf_image = fitz.open(
                "pdf",
                pdf_bytes,
            )

            try:
                document.insert_pdf(pdf_image)
            finally:
                pdf_image.close()

        if document.page_count == 0:
            raise ErreurConversion(
                "Aucune page n'a été créée à partir des images."
            )

        document.save(destination_path)

    except ErreurConversion:
        raise

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de regrouper les images en PDF : {erreur}"
        ) from erreur

    finally:
        document.close()

    return ResultatConversion(
        source=premiere_source,
        destination=destination_path,
        type_conversion="Images → PDF",
    )


def fusionner_pdfs(
    sources: list[str | Path] | tuple[str | Path, ...],
    destination: str | Path | None = None,
) -> ResultatConversion:
    # Fusionne plusieurs PDF dans l'ordre fourni.
    if not sources:
        raise ValueError(
            "Sélectionnez au moins un fichier PDF à fusionner."
        )

    chemins = [
        _verifier_source(source, {".pdf"})
        for source in sources
    ]

    premiere_source = chemins[0]

    if destination is None:
        destination_path = premiere_source.with_name(
            f"{premiere_source.stem}_fusion.pdf"
        )
    else:
        destination_path = Path(destination)

    if destination_path.suffix.lower() != ".pdf":
        raise ValueError(
            "Le fichier de destination doit avoir l'extension .pdf."
        )

    destination_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = fitz.open()

    try:
        for chemin_pdf in chemins:
            pdf = fitz.open(chemin_pdf)

            try:
                if pdf.page_count == 0:
                    continue

                document.insert_pdf(pdf)

            finally:
                pdf.close()

        if document.page_count == 0:
            raise ErreurConversion(
                "Aucune page PDF valide n'a été trouvée."
            )

        document.save(destination_path)

    except ErreurConversion:
        raise

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de fusionner les PDF : {erreur}"
        ) from erreur

    finally:
        document.close()

    return ResultatConversion(
        source=premiere_source,
        destination=destination_path,
        type_conversion="PDFs → PDF",
    )


def _normaliser_pages_selection(
    selection: str,
    nombre_pages: int,
) -> list[int]:
    # Convertit "1,3-5" en indices 0-based ordonnés et uniques.
    texte = selection.strip()

    if not texte:
        return list(range(nombre_pages))

    pages: list[int] = []

    for morceau in texte.split(","):
        morceau = morceau.strip()
        if not morceau:
            continue

        if "-" in morceau:
            debut_txt, fin_txt = morceau.split("-", 1)

            if not debut_txt.strip().isdigit() or not fin_txt.strip().isdigit():
                raise ValueError(
                    "Sélection de pages invalide. Exemple accepté : 1,3-5"
                )

            debut = int(debut_txt)
            fin = int(fin_txt)

            if debut > fin:
                raise ValueError(
                    "Une plage de pages doit aller du plus petit au plus grand."
                )

            numeros = range(debut, fin + 1)

        else:
            if not morceau.isdigit():
                raise ValueError(
                    "Sélection de pages invalide. Exemple accepté : 1,3-5"
                )

            numeros = [int(morceau)]

        for numero in numeros:
            if numero < 1 or numero > nombre_pages:
                raise ValueError(
                    f"La page {numero} n'existe pas. "
                    f"Le document contient {nombre_pages} page(s)."
                )

            index = numero - 1

            if index not in pages:
                pages.append(index)

    if not pages:
        raise ValueError(
            "Aucune page valide n'a été sélectionnée."
        )

    return pages


def pdf_vers_images(
    source: str | Path,
    dossier_destination: str | Path | None = None,
    *,
    format_image: str = "png",
    dpi: int = 150,
    pages: str = "",
) -> list[Path]:
    # Convertit les pages PDF choisies en images séparées.
    source_path = _verifier_source(source, {".pdf"})

    format_normalise = format_image.strip().lower()
    if format_normalise == "jpeg":
        format_normalise = "jpg"

    if format_normalise not in {"png", "jpg"}:
        raise ValueError(
            "Le format image doit être 'png' ou 'jpg'."
        )

    if dpi < 72 or dpi > 600:
        raise ValueError(
            "Le DPI doit être compris entre 72 et 600."
        )

    if dossier_destination is None:
        destination_dir = source_path.with_name(
            f"{source_path.stem}_images"
        )
    else:
        destination_dir = Path(dossier_destination)

    destination_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    zoom = dpi / 72
    matrice = fitz.Matrix(zoom, zoom)
    sorties: list[Path] = []

    try:
        with fitz.open(source_path) as document:
            if document.page_count == 0:
                raise ErreurConversion(
                    "Le PDF ne contient aucune page."
                )

            pages_selectionnees = _normaliser_pages_selection(
                pages,
                document.page_count,
            )

            largeur_numero = max(
                3,
                len(str(document.page_count)),
            )

            for index in pages_selectionnees:
                page = document[index]

                pixmap = page.get_pixmap(
                    matrix=matrice,
                    alpha=False,
                )

                destination = destination_dir / (
                    f"{source_path.stem}_page_"
                    f"{index + 1:0{largeur_numero}d}."
                    f"{format_normalise}"
                )

                pixmap.save(destination)
                sorties.append(destination)

    except (ErreurConversion, ValueError):
        raise

    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de convertir le PDF en images : {erreur}"
        ) from erreur

    return sorties


def csv_vers_excel(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    source_path = _verifier_source(source, {".csv"})
    destination_path = _preparer_destination(
        source_path, destination, ".xlsx"
    )
    from openpyxl import Workbook

    with source_path.open("r", encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        rows = list(csv.reader(f, dialect))

    if not rows:
        raise ErreurConversion("Le fichier CSV est vide.")

    wb = Workbook()
    ws = wb.active
    ws.title = "Données"
    for row in rows:
        ws.append(row)
    for cell in ws[1]:
        cell.font = cell.font.copy(bold=True)
    wb.save(destination_path)

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="CSV → Excel",
    )


def _extraire_tableaux_pdf(source_path: Path):
    """Extrait les tableaux natifs détectables dans un PDF."""
    document = fitz.open(source_path)
    resultats = []

    try:
        for page_no, page in enumerate(document, start=1):
            detection = page.find_tables()

            for table_no, table in enumerate(
                detection.tables,
                start=1,
            ):
                data = table.extract()

                if data and any(
                    any(v not in (None, "") for v in row)
                    for row in data
                ):
                    resultats.append(
                        (page_no, table_no, data)
                    )
    finally:
        document.close()

    return resultats


def _extraire_ocr_pages_pdf(
    source_path: Path,
) -> dict[int, str]:
    """OCR local uniquement des pages PDF qui en ont besoin."""
    from tempfile import TemporaryDirectory

    from .ocr_extractor import extraire_texte_image
    from .ocr_normalizer import normaliser_montants_ocr
    from .pdf_type_detector import analyser_pdf

    analyse = analyser_pdf(source_path)

    if not analyse.necessite_ocr:
        return {}

    pages_ocr = set(analyse.pages_ocr)
    textes: dict[int, str] = {}

    with fitz.open(source_path) as document:
        with TemporaryDirectory(
            prefix="comptaprivee_conversion_ocr_"
        ) as dossier_temporaire:
            dossier = Path(dossier_temporaire)

            for numero_page, page in enumerate(
                document,
                start=1,
            ):
                if numero_page not in pages_ocr:
                    continue

                image_path = dossier / f"page_{numero_page}.png"

                pixmap = page.get_pixmap(
                    dpi=300,
                    alpha=False,
                )
                pixmap.save(image_path)

                texte = normaliser_montants_ocr(
                    extraire_texte_image(
                        image_path
                    ).strip()
                )

                if texte:
                    textes[numero_page] = texte

    return textes


def _extraire_pages_pdf_avec_ocr(
    source: str | Path,
) -> list[str]:
    # Réutilise l'extracteur PDF intelligent déjà validé par ComptaPrivée AI
    # puis restitue une liste page par page pour les fallbacks CSV / Excel.
    from .pdf_extractor import extraire_texte_pdf

    texte = extraire_texte_pdf(source)

    if not texte.strip():
        return [""]

    separateur = "\n\n--- Page suivante ---\n\n"

    return [
        bloc.strip()
        for bloc in texte.split(separateur)
    ]


def _extraire_pages_pdf_fallback(
    source_path: Path,
) -> dict[int, str]:
    # Combine le texte natif de chaque page et l'OCR uniquement
    # pour les pages qui en ont besoin.
    pages: dict[int, str] = {}

    with fitz.open(source_path) as document:
        for numero_page, page in enumerate(document, start=1):
            texte = page.get_text("text").strip()
            if texte:
                pages[numero_page] = texte

    # L'ancien moteur OCR reste la source prioritaire pour les pages
    # scannées / image. Les tests historiques peuvent aussi le monkeypatcher.
    pages_ocr = _extraire_ocr_pages_pdf(source_path)

    for numero_page, texte_ocr in pages_ocr.items():
        if texte_ocr and texte_ocr.strip():
            pages[numero_page] = texte_ocr.strip()

    return pages


def pdf_vers_csv(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    source_path = _verifier_source(source, {".pdf"})
    destination_path = _preparer_destination(
        source_path, destination, ".csv"
    )

    tableaux = extraire_tableaux_pdf(source_path)

    with destination_path.open(
        "w",
        encoding="utf-8-sig",
        newline="",
    ) as fichier_csv:
        writer = csv.writer(fichier_csv)

        if tableaux:
            plusieurs = len(tableaux) > 1

            for index, tableau in enumerate(tableaux, start=1):
                if plusieurs:
                    writer.writerow(
                        [
                            f"Page {tableau.numero_page}",
                            f"Tableau {tableau.numero_tableau}",
                        ]
                    )

                for ligne in _preparer_tableau_ocr_pour_csv(tableau.lignes):
                    writer.writerow(ligne)

                if plusieurs and index < len(tableaux):
                    writer.writerow([])

        else:
            pages_ocr = _extraire_pages_pdf_fallback(source_path)
            writer.writerow(["Page", "Ligne", "Texte OCR"])

            for numero_page, texte in sorted(pages_ocr.items()):
                lignes = [
                    ligne.strip()
                    for ligne in texte.splitlines()
                    if ligne.strip()
                ]

                for numero_ligne, ligne in enumerate(lignes, start=1):
                    writer.writerow(
                        [numero_page, numero_ligne, ligne]
                    )

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="PDF → CSV",
    )


def _preparer_tableau_ocr_pour_excel(
    lignes: list[list[str]],
) -> list[list[str]]:
    """Prépare un tableau OCR pour Excel sans inventer de valeurs.

    - normalise seulement les montants sûrs ;
    - ajoute une colonne « Validation OCR » ;
    - marque « À VÉRIFIER » si les montants sont incohérents ;
    - laisse fournisseur et numéro de facture exactement tels que lus.
    """
    from .ocr_table_validator import (
        detecter_anomalies_comptables,
        normaliser_tableau_comptable,
    )

    if not lignes:
        return lignes

    normalise = normaliser_tableau_comptable(
        lignes,
    )

    anomalies = detecter_anomalies_comptables(
        normalise,
    )

    lignes_a_verifier = {
        anomalie.numero_ligne
        for anomalie in anomalies
    }

    entete = list(normalise[0])

    if not any(
        cellule.strip().lower() == "validation ocr"
        for cellule in entete
    ):
        entete.append("Validation OCR")

    resultat = [entete]

    for numero_ligne, ligne in enumerate(
        normalise[1:],
        start=2,
    ):
        copie = list(ligne)

        statut = (
            "À VÉRIFIER"
            if numero_ligne in lignes_a_verifier
            else "OK"
        )

        copie.append(statut)
        resultat.append(copie)

    return resultat


def _preparer_tableau_ocr_pour_csv(
    lignes: list[list[str]],
) -> list[list[str]]:
    """Prépare un tableau OCR pour CSV avec les mêmes contrôles que l'Excel.

    Réutilise volontairement la logique validée de PDF -> Excel afin que
    les deux formats aient exactement le même comportement :
    normalisation prudente + colonne Validation OCR.
    """
    return _preparer_tableau_ocr_pour_excel(
        lignes,
    )


def pdf_vers_excel(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    source_path = _verifier_source(source, {".pdf"})
    destination_path = _preparer_destination(
        source_path, destination, ".xlsx"
    )

    from openpyxl import Workbook
    from openpyxl.styles import Font

    tableaux = extraire_tableaux_pdf(source_path)

    wb = Workbook()
    ws = wb.active

    if tableaux:
        wb.remove(ws)

        for tableau in tableaux:
            titre = f"P{tableau.numero_page}_T{tableau.numero_tableau}"
            ws = wb.create_sheet(title=titre[:31])

            for ligne in _preparer_tableau_ocr_pour_excel(tableau.lignes):
                ws.append(ligne)

            if ws.max_row >= 1:
                for cell in ws[1]:
                    cell.font = Font(bold=True)

            for colonne in ws.columns:
                largeur = max(
                    len(str(cell.value)) if cell.value is not None else 0
                    for cell in colonne
                )
                ws.column_dimensions[colonne[0].column_letter].width = min(
                    max(largeur + 2, 10),
                    45,
                )

    else:
        pages_ocr = _extraire_pages_pdf_fallback(source_path)

        wb.remove(ws)

        if not pages_ocr:
            ws = wb.create_sheet(title="OCR_Page_1")
            ws.append(["Ligne", "Texte OCR"])
        else:
            for numero_page, texte in sorted(pages_ocr.items()):
                ws = wb.create_sheet(
                    title=f"OCR_Page_{numero_page}"[:31]
                )
                ws.append(["Page", "Ligne", "Texte OCR"])

                lignes = [
                    ligne.strip()
                    for ligne in texte.splitlines()
                    if ligne.strip()
                ]

                for numero_ligne, ligne in enumerate(lignes, start=1):
                    ws.append(
                        [numero_page, numero_ligne, ligne]
                    )

                for cell in ws[1]:
                    cell.font = Font(bold=True)

                ws.column_dimensions["A"].width = 10
                ws.column_dimensions["B"].width = 10
                ws.column_dimensions["C"].width = 80

    wb.save(destination_path)

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="PDF → Excel",
    )


def pdf_vers_word(
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Convertit PDF natif/mixte/scanné en Word avec OCR local."""
    source_path = _verifier_source(
        source,
        {".pdf"},
    )
    destination_path = _preparer_destination(
        source_path,
        destination,
        ".docx",
    )

    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as erreur:
        raise ErreurConversion(
            "PDF → Word nécessite python-docx."
        ) from erreur

    textes_ocr = _extraire_ocr_pages_pdf(
        source_path
    )
    pdf = fitz.open(source_path)

    try:
        document_word = Document()
        contenu_trouve = False

        for index_page, page in enumerate(pdf):
            numero_page = index_page + 1

            if index_page > 0:
                document_word.add_page_break()

            detection = page.find_tables()
            zones_tableaux = []
            tableaux = []

            for tableau in detection.tables:
                donnees = tableau.extract()

                if not donnees:
                    continue

                if not any(
                    any(v not in (None, "") for v in ligne)
                    for ligne in donnees
                ):
                    continue

                zones_tableaux.append(
                    fitz.Rect(tableau.bbox)
                )
                tableaux.append(donnees)

            blocs = page.get_text(
                "blocks",
                sort=True,
            )

            for bloc in blocs:
                rect_bloc = fitz.Rect(
                    bloc[0],
                    bloc[1],
                    bloc[2],
                    bloc[3],
                )

                if any(
                    rect_bloc.intersects(zone)
                    for zone in zones_tableaux
                ):
                    continue

                texte = str(bloc[4]).strip()

                if not texte:
                    continue

                contenu_trouve = True

                for ligne in texte.splitlines():
                    ligne = ligne.strip()

                    if ligne:
                        document_word.add_paragraph(
                            ligne
                        )

            for donnees in tableaux:
                nb_colonnes = max(
                    len(ligne)
                    for ligne in donnees
                )

                tableau_word = document_word.add_table(
                    rows=0,
                    cols=nb_colonnes,
                )
                tableau_word.style = "Table Grid"
                tableau_word.alignment = (
                    WD_TABLE_ALIGNMENT.CENTER
                )

                for numero_ligne, ligne in enumerate(
                    donnees
                ):
                    cellules = tableau_word.add_row().cells

                    for numero_colonne in range(
                        nb_colonnes
                    ):
                        valeur = (
                            ligne[numero_colonne]
                            if numero_colonne < len(ligne)
                            else ""
                        )
                        cellules[
                            numero_colonne
                        ].text = (
                            ""
                            if valeur is None
                            else str(valeur)
                        )

                        for paragraphe in cellules[
                            numero_colonne
                        ].paragraphs:
                            paragraphe.alignment = (
                                WD_ALIGN_PARAGRAPH.LEFT
                            )

                            if numero_ligne == 0:
                                for run in paragraphe.runs:
                                    run.bold = True

                contenu_trouve = True

            # Si la page est un scan, ajouter son texte OCR.
            texte_ocr = textes_ocr.get(
                numero_page,
                "",
            ).strip()

            if texte_ocr:
                document_word.add_paragraph(
                    "Texte reconnu par OCR"
                ).runs[0].bold = True

                for ligne in texte_ocr.splitlines():
                    ligne = ligne.strip()

                    if ligne:
                        document_word.add_paragraph(
                            ligne
                        )

                contenu_trouve = True

        if not contenu_trouve:
            raise ErreurConversion(
                "Aucun texte, tableau ou résultat OCR "
                "n'a été détecté dans ce PDF."
            )

        document_word.save(
            destination_path
        )

    except ErreurConversion:
        raise
    except Exception as erreur:
        raise ErreurConversion(
            f"Impossible de convertir PDF en Word : {erreur}"
        ) from erreur
    finally:
        pdf.close()

    return ResultatConversion(
        source=source_path,
        destination=destination_path,
        type_conversion="PDF → Word",
    )


def convertir_document(
    type_conversion: str,
    source: str | Path,
    destination: str | Path | None = None,
) -> ResultatConversion:
    """Route une conversion vers le moteur correspondant."""
    conversions = {
        "Word → PDF": word_vers_pdf,
        "Excel → PDF": excel_vers_pdf,
        "Excel → CSV": excel_vers_csv,
        "CSV → Excel": csv_vers_excel,
        "Image → PDF": image_vers_pdf,
        "PDF → CSV": pdf_vers_csv,
        "PDF → Excel": pdf_vers_excel,
        "PDF → Word": pdf_vers_word,
    }

    fonction = conversions.get(type_conversion)

    if fonction is None:
        raise ValueError(
            f"Conversion non prise en charge : {type_conversion}"
        )

    return fonction(
        source,
        destination,
    )
