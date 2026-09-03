"""Tests du moteur de conversion locale."""

from pathlib import Path

import pytest

from src.comptaprivee.document_converter import (
    CONVERSIONS_SUPPORTEES,
    convertir_document,
    image_vers_pdf,
)


def test_liste_conversions_supportees() -> None:
    assert "Word → PDF" in CONVERSIONS_SUPPORTEES
    assert "Excel → PDF" in CONVERSIONS_SUPPORTEES
    assert "Excel → CSV" in CONVERSIONS_SUPPORTEES
    assert "Images → PDF" in CONVERSIONS_SUPPORTEES


def test_refuser_conversion_inconnue(tmp_path) -> None:
    source = tmp_path / "test.txt"
    source.write_text("demo", encoding="utf-8")

    with pytest.raises(ValueError):
        convertir_document(
            "TXT → MP3",
            source,
        )


def test_refuser_source_inexistante(tmp_path) -> None:
    with pytest.raises(FileNotFoundError):
        image_vers_pdf(
            tmp_path / "absente.png",
        )


def test_refuser_mauvaise_extension_image(tmp_path) -> None:
    source = tmp_path / "document.txt"
    source.write_text("demo", encoding="utf-8")

    with pytest.raises(ValueError):
        image_vers_pdf(source)


def test_image_vers_pdf(tmp_path) -> None:
    import fitz

    source = tmp_path / "image.png"
    destination = tmp_path / "image_convertie.pdf"

    document = fitz.open()
    page = document.new_page(width=200, height=120)
    pixmap = page.get_pixmap()
    pixmap.save(source)
    document.close()

    resultat = image_vers_pdf(
        source,
        destination,
    )

    assert resultat.destination == destination
    assert destination.exists()

    pdf = fitz.open(destination)
    try:
        assert pdf.page_count == 1
    finally:
        pdf.close()


def test_destination_pdf_obligatoire(tmp_path) -> None:
    import fitz

    source = tmp_path / "image.png"
    mauvaise_destination = tmp_path / "sortie.txt"

    document = fitz.open()
    page = document.new_page(width=50, height=50)
    page.get_pixmap().save(source)
    document.close()

    with pytest.raises(ValueError):
        image_vers_pdf(
            source,
            mauvaise_destination,
        )

def _creer_pdf_tableau_test(chemin) -> None:
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=400, height=260)
    xs = [40, 150, 280, 360]
    ys = [40, 70, 100, 130]
    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))
    data = [
        ["Date", "Fournisseur", "Total"],
        ["2026-09-01", "A", "114.98"],
        ["2026-09-02", "B", "287.44"],
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            page.insert_text((xs[c] + 4, ys[r] + 20), value, fontsize=8)
    doc.save(chemin)
    doc.close()


def test_conversions_avancees_dans_liste() -> None:
    from src.comptaprivee.document_converter import CONVERSIONS_SUPPORTEES
    for nom in ("CSV → Excel", "PDF → CSV", "PDF → Excel", "PDF → Word"):
        assert nom in CONVERSIONS_SUPPORTEES


def test_csv_vers_excel(tmp_path) -> None:
    from openpyxl import load_workbook
    from src.comptaprivee.document_converter import csv_vers_excel
    src = tmp_path / "a.csv"
    dst = tmp_path / "a.xlsx"
    src.write_text("Date,Fournisseur,Total\n2026-09-01,A,114.98\n", encoding="utf-8")
    csv_vers_excel(src, dst)
    wb = load_workbook(dst)
    try:
        ws = wb.active
        assert ws["A1"].value == "Date"
        assert ws["B2"].value == "A"
    finally:
        wb.close()


def test_pdf_vers_csv(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_csv
    src = tmp_path / "t.pdf"
    dst = tmp_path / "t.csv"
    _creer_pdf_tableau_test(src)
    pdf_vers_csv(src, dst)
    contenu = dst.read_text(encoding="utf-8-sig")
    assert "Fournisseur" in contenu
    assert "114.98" in contenu


def test_pdf_vers_excel(tmp_path) -> None:
    from openpyxl import load_workbook
    from src.comptaprivee.document_converter import pdf_vers_excel
    src = tmp_path / "t.pdf"
    dst = tmp_path / "t.xlsx"
    _creer_pdf_tableau_test(src)
    pdf_vers_excel(src, dst)
    wb = load_workbook(dst)
    try:
        ws = wb.worksheets[0]
        assert ws["A1"].value == "Date"
        assert ws["C2"].value == "114.98"
    finally:
        wb.close()


def test_pdf_vers_word(tmp_path) -> None:
    import fitz
    from docx import Document
    from src.comptaprivee.document_converter import pdf_vers_word
    src = tmp_path / "t.pdf"
    dst = tmp_path / "t.docx"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Facture FAC-001 - Total 114.98 CAD")
    pdf.save(src)
    pdf.close()
    pdf_vers_word(src, dst)
    doc = Document(dst)
    texte = "\n".join(p.text for p in doc.paragraphs)
    assert "FAC-001" in texte

def test_pdf_vers_word_recree_tableau(tmp_path) -> None:
    import fitz
    from docx import Document
    from src.comptaprivee.document_converter import pdf_vers_word

    source = tmp_path / "tableau_word.pdf"
    destination = tmp_path / "tableau_word.docx"

    pdf = fitz.open()
    page = pdf.new_page(width=400, height=260)

    xs = [40, 150, 280, 360]
    ys = [40, 70, 100, 130]

    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))

    data = [
        ["Date", "Fournisseur", "Total"],
        ["2026-09-01", "A", "114.98"],
        ["2026-09-02", "B", "287.44"],
    ]

    for r, row in enumerate(data):
        for c, value in enumerate(row):
            page.insert_text(
                (xs[c] + 4, ys[r] + 20),
                value,
                fontsize=8,
            )

    pdf.save(source)
    pdf.close()

    pdf_vers_word(source, destination)

    doc = Document(destination)

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Date"
    assert doc.tables[0].cell(1, 1).text == "A"
    assert doc.tables[0].cell(2, 2).text == "287.44"


def test_pdf_vers_word_evite_dupliquer_texte_tableau(tmp_path) -> None:
    import fitz
    from docx import Document
    from src.comptaprivee.document_converter import pdf_vers_word

    source = tmp_path / "mixte.pdf"
    destination = tmp_path / "mixte.docx"

    pdf = fitz.open()
    page = pdf.new_page(width=400, height=300)
    page.insert_text((40, 30), "Rapport comptable")

    xs = [40, 180, 320]
    ys = [70, 100, 130]

    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))

    page.insert_text((45, 90), "Facture", fontsize=8)
    page.insert_text((185, 90), "Total", fontsize=8)
    page.insert_text((45, 120), "FAC-001", fontsize=8)
    page.insert_text((185, 120), "114.98", fontsize=8)

    pdf.save(source)
    pdf.close()

    pdf_vers_word(source, destination)

    doc = Document(destination)
    texte_paragraphes = "\n".join(
        p.text for p in doc.paragraphs
    )

    assert "Rapport comptable" in texte_paragraphes
    assert len(doc.tables) == 1
    assert "FAC-001" not in texte_paragraphes

def test_pdf_vers_word_fallback_ocr(
    tmp_path,
    monkeypatch,
) -> None:
    import fitz
    from docx import Document
    import src.comptaprivee.document_converter as dc

    source = tmp_path / "scan.pdf"
    destination = tmp_path / "scan.docx"

    doc = fitz.open()
    doc.new_page()
    doc.save(source)
    doc.close()

    monkeypatch.setattr(
        dc,
        "_extraire_ocr_pages_pdf",
        lambda _source: {
            1: "Facture FAC-OCR-001\nTotal 114.98 CAD"
        },
    )

    dc.pdf_vers_word(
        source,
        destination,
    )

    word = Document(destination)
    texte = "\n".join(
        p.text for p in word.paragraphs
    )

    assert "FAC-OCR-001" in texte
    assert "114.98" in texte


def test_pdf_vers_csv_fallback_ocr(
    tmp_path,
    monkeypatch,
) -> None:
    import fitz
    import src.comptaprivee.document_converter as dc

    source = tmp_path / "scan.pdf"
    destination = tmp_path / "scan.csv"

    doc = fitz.open()
    doc.new_page()
    doc.save(source)
    doc.close()

    monkeypatch.setattr(
        dc,
        "_extraire_ocr_pages_pdf",
        lambda _source: {
            1: "Facture FAC-OCR-001\nTotal 114.98 CAD"
        },
    )

    dc.pdf_vers_csv(
        source,
        destination,
    )

    contenu = destination.read_text(
        encoding="utf-8-sig"
    )

    assert "Texte OCR" in contenu
    assert "FAC-OCR-001" in contenu
    assert "114.98" in contenu


def test_pdf_vers_excel_fallback_ocr(
    tmp_path,
    monkeypatch,
) -> None:
    import fitz
    from openpyxl import load_workbook
    import src.comptaprivee.document_converter as dc

    source = tmp_path / "scan.pdf"
    destination = tmp_path / "scan.xlsx"

    doc = fitz.open()
    doc.new_page()
    doc.save(source)
    doc.close()

    monkeypatch.setattr(
        dc,
        "_extraire_ocr_pages_pdf",
        lambda _source: {
            1: "Facture FAC-OCR-001\nTotal 114.98 CAD"
        },
    )

    dc.pdf_vers_excel(
        source,
        destination,
    )

    wb = load_workbook(destination)
    try:
        assert "OCR_Page_1" in wb.sheetnames
        ws = wb["OCR_Page_1"]
        valeurs = [
            str(ws.cell(row=i, column=3).value or "")
            for i in range(1, ws.max_row + 1)
        ]
        texte = "\n".join(valeurs)
        assert "FAC-OCR-001" in texte
        assert "114.98" in texte
    finally:
        wb.close()

def _creer_png_test(chemin, texte: str) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=300, height=180)
    page.insert_text((30, 90), texte)
    pix = page.get_pixmap()
    pix.save(chemin)
    doc.close()


def test_images_vers_pdf_plusieurs_pages(tmp_path) -> None:
    import fitz

    from src.comptaprivee.document_converter import images_vers_pdf

    image1 = tmp_path / "page1.png"
    image2 = tmp_path / "page2.png"
    destination = tmp_path / "regroupe.pdf"

    _creer_png_test(image1, "IMAGE 1")
    _creer_png_test(image2, "IMAGE 2")

    resultat = images_vers_pdf(
        [image1, image2],
        destination,
    )

    assert resultat.type_conversion == "Images → PDF"
    assert resultat.destination == destination
    assert destination.exists()

    with fitz.open(destination) as pdf:
        assert pdf.page_count == 2


def test_images_vers_pdf_destination_automatique(tmp_path) -> None:
    from src.comptaprivee.document_converter import images_vers_pdf

    image1 = tmp_path / "piece_1.png"
    image2 = tmp_path / "piece_2.png"

    _creer_png_test(image1, "A")
    _creer_png_test(image2, "B")

    resultat = images_vers_pdf([image1, image2])

    assert resultat.destination.name == "piece_1_images.pdf"
    assert resultat.destination.exists()


def test_images_vers_pdf_refuse_liste_vide() -> None:
    import pytest

    from src.comptaprivee.document_converter import images_vers_pdf

    with pytest.raises(ValueError):
        images_vers_pdf([])


def test_images_vers_pdf_refuse_format_non_image(tmp_path) -> None:
    import pytest

    from src.comptaprivee.document_converter import images_vers_pdf

    fichier = tmp_path / "notes.txt"
    fichier.write_text("pas une image", encoding="utf-8")

    with pytest.raises(ValueError):
        images_vers_pdf([fichier])

def test_menu_conversion_ne_propose_qu_images_vers_pdf() -> None:
    from src.comptaprivee.document_converter import CONVERSIONS_SUPPORTEES

    assert "Images → PDF" in CONVERSIONS_SUPPORTEES
    assert "Image → PDF" not in CONVERSIONS_SUPPORTEES

def _creer_pdf_test(chemin, texte: str, pages: int = 1) -> None:
    import fitz

    doc = fitz.open()

    for index in range(pages):
        page = doc.new_page()
        page.insert_text(
            (72, 72),
            f"{texte} - page {index + 1}",
        )

    doc.save(chemin)
    doc.close()


def test_fusionner_pdfs_conserve_ordre_et_pages(tmp_path) -> None:
    import fitz

    from src.comptaprivee.document_converter import fusionner_pdfs

    pdf1 = tmp_path / "a.pdf"
    pdf2 = tmp_path / "b.pdf"
    destination = tmp_path / "fusion.pdf"

    _creer_pdf_test(pdf1, "PDF A", pages=2)
    _creer_pdf_test(pdf2, "PDF B", pages=1)

    resultat = fusionner_pdfs(
        [pdf1, pdf2],
        destination,
    )

    assert resultat.type_conversion == "PDFs → PDF"
    assert destination.exists()

    with fitz.open(destination) as pdf:
        assert pdf.page_count == 3
        assert "PDF A" in pdf[0].get_text()
        assert "PDF A" in pdf[1].get_text()
        assert "PDF B" in pdf[2].get_text()


def test_fusionner_pdfs_destination_automatique(tmp_path) -> None:
    from src.comptaprivee.document_converter import fusionner_pdfs

    pdf1 = tmp_path / "facture_1.pdf"
    pdf2 = tmp_path / "facture_2.pdf"

    _creer_pdf_test(pdf1, "A")
    _creer_pdf_test(pdf2, "B")

    resultat = fusionner_pdfs([pdf1, pdf2])

    assert resultat.destination.name == "facture_1_fusion.pdf"
    assert resultat.destination.exists()


def test_fusionner_pdfs_refuse_liste_vide() -> None:
    import pytest

    from src.comptaprivee.document_converter import fusionner_pdfs

    with pytest.raises(ValueError):
        fusionner_pdfs([])


def test_fusionner_pdfs_refuse_format_non_pdf(tmp_path) -> None:
    import pytest

    from src.comptaprivee.document_converter import fusionner_pdfs

    fichier = tmp_path / "notes.txt"
    fichier.write_text("pas un pdf", encoding="utf-8")

    with pytest.raises(ValueError):
        fusionner_pdfs([fichier])

def test_pdf_vers_images_png_une_image_par_page(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "PDF TEST", pages=3)

    dossier = tmp_path / "images"
    sorties = pdf_vers_images(
        source,
        dossier,
        format_image="png",
        dpi=120,
    )

    assert len(sorties) == 3
    assert [p.name for p in sorties] == [
        "source_page_001.png",
        "source_page_002.png",
        "source_page_003.png",
    ]
    assert all(p.exists() for p in sorties)


def test_pdf_vers_images_jpg(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "PDF TEST", pages=2)

    sorties = pdf_vers_images(
        source,
        tmp_path / "jpg",
        format_image="jpg",
    )

    assert len(sorties) == 2
    assert sorties[0].suffix == ".jpg"
    assert sorties[1].suffix == ".jpg"


def test_pdf_vers_images_destination_automatique(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "facture.pdf"
    _creer_pdf_test(source, "FACTURE")

    sorties = pdf_vers_images(source)

    assert len(sorties) == 1
    assert sorties[0].parent.name == "facture_images"
    assert sorties[0].exists()


def test_pdf_vers_images_refuse_format_invalide(tmp_path) -> None:
    import pytest

    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "TEST")

    with pytest.raises(ValueError):
        pdf_vers_images(
            source,
            format_image="bmp",
        )


def test_pdf_vers_images_refuse_dpi_invalide(tmp_path) -> None:
    import pytest

    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "TEST")

    with pytest.raises(ValueError):
        pdf_vers_images(
            source,
            dpi=20,
        )

def test_normaliser_pages_selection_plages_et_doublons() -> None:
    from src.comptaprivee.document_converter import _normaliser_pages_selection

    assert _normaliser_pages_selection("1,3-5,3", 6) == [0, 2, 3, 4]


def test_pdf_vers_images_selection_pages(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "PDF TEST", pages=5)

    sorties = pdf_vers_images(
        source,
        tmp_path / "selection",
        pages="2,4-5",
    )

    assert [p.name for p in sorties] == [
        "source_page_002.png",
        "source_page_004.png",
        "source_page_005.png",
    ]


def test_pdf_vers_images_selection_vide_convertit_tout(tmp_path) -> None:
    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "PDF TEST", pages=2)

    sorties = pdf_vers_images(
        source,
        tmp_path / "tout",
        pages="",
    )

    assert len(sorties) == 2


def test_pdf_vers_images_refuse_page_hors_limite(tmp_path) -> None:
    import pytest

    from src.comptaprivee.document_converter import pdf_vers_images

    source = tmp_path / "source.pdf"
    _creer_pdf_test(source, "PDF TEST", pages=2)

    with pytest.raises(ValueError):
        pdf_vers_images(
            source,
            tmp_path / "images",
            pages="3",
        )

def test_pdf_vers_csv_utilise_tableau_detecte(tmp_path) -> None:
    import csv
    import fitz
    from src.comptaprivee.document_converter import pdf_vers_csv

    source = tmp_path / "tableau.pdf"
    doc = fitz.open()
    page = doc.new_page(width=500, height=300)
    xs = [40, 160, 280]
    ys = [40, 80, 120]

    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))

    page.insert_text((45, 65), "Date")
    page.insert_text((165, 65), "Total")
    page.insert_text((45, 105), "2026-09-01")
    page.insert_text((165, 105), "114.98")

    doc.save(source)
    doc.close()

    destination = tmp_path / "sortie.csv"
    pdf_vers_csv(source, destination)

    with destination.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))

    assert rows[0][:2] == ["Date", "Total"]
    assert rows[1][:2] == ["2026-09-01", "114.98"]


def test_pdf_vers_excel_cree_une_feuille_par_tableau(tmp_path) -> None:
    import fitz
    from openpyxl import load_workbook
    from src.comptaprivee.document_converter import pdf_vers_excel

    source = tmp_path / "tableau.pdf"
    doc = fitz.open()
    page = doc.new_page(width=500, height=300)
    xs = [40, 160, 280]
    ys = [40, 80, 120]

    for x in xs:
        page.draw_line((x, ys[0]), (x, ys[-1]))
    for y in ys:
        page.draw_line((xs[0], y), (xs[-1], y))

    page.insert_text((45, 65), "Fournisseur")
    page.insert_text((165, 65), "Total")
    page.insert_text((45, 105), "ABC Inc.")
    page.insert_text((165, 105), "287.44")

    doc.save(source)
    doc.close()

    destination = tmp_path / "sortie.xlsx"
    pdf_vers_excel(source, destination)

    wb = load_workbook(destination)
    ws = wb[wb.sheetnames[0]]
    assert ws["A1"].value == "Fournisseur"
    assert ws["B1"].value == "Total"
    assert ws["A2"].value == "ABC Inc."
    assert ws["B2"].value == "287.44"
    wb.close()


def test_pdf_vers_csv_sans_tableau_garde_fallback_ocr(tmp_path) -> None:
    import csv
    import fitz
    from src.comptaprivee.document_converter import pdf_vers_csv

    source = tmp_path / "texte.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Bonjour comptable")
    doc.save(source)
    doc.close()

    destination = tmp_path / "texte.csv"
    pdf_vers_csv(source, destination)

    with destination.open("r", encoding="utf-8-sig", newline="") as f:
        rows = list(csv.reader(f))

    assert rows[0] == ["Page", "Ligne", "Texte OCR"]
    assert rows[1][2] == "Bonjour comptable"


def test_pdf_vers_excel_sans_tableau_garde_feuille_ocr(tmp_path) -> None:
    import fitz
    from openpyxl import load_workbook
    from src.comptaprivee.document_converter import pdf_vers_excel

    source = tmp_path / "texte.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Bonjour comptable")
    doc.save(source)
    doc.close()

    destination = tmp_path / "texte.xlsx"
    pdf_vers_excel(source, destination)

    wb = load_workbook(destination)
    assert "OCR_Page_1" in wb.sheetnames
    assert wb["OCR_Page_1"]["C2"].value == "Bonjour comptable"
    wb.close()

def test_preparer_tableau_ocr_excel_marque_incoherent() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_excel,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-05",
            "Fournisseur F",
            "FAC-N05",
            "1200 00",
            "40 00",
            "119 70",
            "1279 70",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_excel(
        tableau,
    )

    assert resultat[0][-1] == "Validation OCR"
    assert resultat[1][3:7] == [
        "1200.00",
        "40.00",
        "119.70",
        "1279.70",
    ]
    assert resultat[1][-1] == "À VÉRIFIER"


def test_preparer_tableau_ocr_excel_ligne_valide_ok() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_excel,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-01",
            "Fournisseur A",
            "FAC-001",
            "100.00",
            "5.00",
            "9.98",
            "114.98",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_excel(
        tableau,
    )

    assert resultat[1][-1] == "OK"


def test_preparer_tableau_ocr_excel_ne_devine_pas_identifiants() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_excel,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-05",
            "Fournisseur F",
            "FAC-N05",
            "1200.00",
            "60.00",
            "119.70",
            "1379.70",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_excel(
        tableau,
    )

    assert resultat[1][1] == "Fournisseur F"
    assert resultat[1][2] == "FAC-N05"

def test_preparer_tableau_ocr_csv_marque_incoherent() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_csv,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-05",
            "Fournisseur F",
            "FAC-N05",
            "1200 00",
            "40 00",
            "119 70",
            "1279 70",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_csv(
        tableau,
    )

    assert resultat[0][-1] == "Validation OCR"
    assert resultat[1][3:7] == [
        "1200.00",
        "40.00",
        "119.70",
        "1279.70",
    ]
    assert resultat[1][-1] == "À VÉRIFIER"


def test_preparer_tableau_ocr_csv_ligne_valide_ok() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_csv,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-01",
            "Fournisseur A",
            "FAC-001",
            "100.00",
            "5.00",
            "9.98",
            "114.98",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_csv(
        tableau,
    )

    assert resultat[1][-1] == "OK"


def test_preparer_tableau_ocr_csv_ne_devine_pas_identifiants() -> None:
    from src.comptaprivee.document_converter import (
        _preparer_tableau_ocr_pour_csv,
    )

    tableau = [
        [
            "Date",
            "Fournisseur",
            "No facture",
            "Sous-total",
            "TPS",
            "TVQ",
            "Total",
        ],
        [
            "2026-09-05",
            "Fournisseur F",
            "FAC-N05",
            "1200.00",
            "60.00",
            "119.70",
            "1379.70",
        ],
    ]

    resultat = _preparer_tableau_ocr_pour_csv(
        tableau,
    )

    assert resultat[1][1] == "Fournisseur F"
    assert resultat[1][2] == "FAC-N05"
